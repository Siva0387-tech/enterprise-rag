from pathlib import Path
from datetime import datetime
from typing import List
from utils.metadata import create_metadata
from docx import Document as DocxDocument
from langchain_core.documents import Document


class DOCXLoader:
    """
    Loads DOCX files and returns a list of LangChain Document objects.
    """

    def load(self, file_path: str) -> List[Document]:

        doc = DocxDocument(file_path)

        text = "\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        )

        metadata = create_metadata(
            file_path=file_path,
            document_type="docx"
        )

        return [
            Document(
                page_content=text,
                metadata=metadata
            )
        ]